
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helper functions ──────────────────────────────────────────────────────────
def heading(text, level=1, color=RGBColor(0x2E, 0x86, 0xC1)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = color
        run.font.bold = True
    return p

def para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def code_block(code_text):
    """Add a shaded code block."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    return p

def add_image_if_exists(path, caption="", width=Inches(5)):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(9)

def table_2col(rows):
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    hdr[0].text = 'Item'
    hdr[1].text = 'Details'
    for k, v in rows:
        row = t.add_row().cells
        row[0].text = k
        row[1].text = v
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_heading('Photo Converter — Creative Studio', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1A, 0x5C, 0x99)
    run.font.size = Pt(28)

sub = doc.add_paragraph('AI-Powered Image Transformation Web Application')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(14)
sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('Technology: Python · FastAPI · React · OpenCV · Vite\n').bold = True
info.add_run('Deployment: Render (Backend) · Vercel (Frontend)\n')
info.add_run('Version: 1.0  |  Year: 2025')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. PROJECT OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading('1. Project Overview', 1)
para(
    'Photo Converter — Creative Studio is a full-stack web application that transforms '
    'ordinary photographs into artistic styles using computer vision techniques. '
    'Users can upload any JPG or PNG image and instantly apply one of four creative filters:\n'
    '  • Pencil Sketch\n'
    '  • Anime Art\n'
    '  • Oil Painting\n'
    '  • Background Removal\n\n'
    'The application is deployed live and accessible from any device via a browser. '
    'It was built as a college/personal project to demonstrate real-world skills in '
    'AI, web development, and cloud deployment.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. TECHNOLOGIES USED
# ══════════════════════════════════════════════════════════════════════════════
heading('2. Technologies Used', 1)
table_2col([
    ('Frontend Framework',  'React 18 (with Vite build tool)'),
    ('Backend Framework',   'FastAPI (Python 3.11)'),
    ('Image Processing',    'OpenCV (cv2), NumPy'),
    ('Background Removal',  'rembg + ONNX Runtime (u2netp model)'),
    ('Styling / Animation', 'CSS + Vanta.js HALO (animated 3-D background)'),
    ('Backend Hosting',     'Render.com (free tier, auto-sleep)'),
    ('Frontend Hosting',    'Vercel (free tier, instant CDN)'),
    ('Version Control',     'Git / GitHub'),
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  3. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
heading('3. System Architecture', 1)
para(
    'The application follows a classic client–server architecture:\n\n'
    '  [User Browser]\n'
    '       │  HTTP POST (multipart/form-data)\n'
    '       ▼\n'
    '  [Vercel — React Frontend]\n'
    '       │  fetch() API call to BACKEND_URL\n'
    '       ▼\n'
    '  [Render — FastAPI Backend]\n'
    '       │  OpenCV / rembg processing\n'
    '       ▼\n'
    '  [Processed Image returned as binary response]\n'
    '       │\n'
    '       ▼\n'
    '  [Browser displays result + Download button]\n\n'
    'CORS is fully enabled on the backend so any frontend origin can call the API.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. PROJECT FOLDER STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading('4. Project Folder Structure', 1)
code_block(
"""photo-converter/
│
├── backend/
│   ├── main.py            ← FastAPI application (all image-processing endpoints)
│   ├── requirements.txt   ← Python dependencies
│   └── render.yaml        ← Render deployment configuration
│
├── frontend/
│   ├── index.html         ← HTML entry point (loads Vanta.js)
│   ├── vercel.json        ← Vercel deployment configuration
│   └── src/
│       ├── main.jsx       ← React entry point
│       ├── App.jsx        ← Main React component (UI + API calls)
│       └── index.css      ← Global styles
│
├── start.bat              ← One-click local startup script (Windows)
├── test_anime.py          ← Local test script for anime filter
├── test_anime2.py         ← Improved anime filter test
└── .gitignore
"""
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. STEP-BY-STEP HOW THE APP WORKS
# ══════════════════════════════════════════════════════════════════════════════
heading('5. Step-by-Step: How the App Works', 1)

heading('Step 1 — Open the Web App', 2)
para('The user opens the deployed URL in any browser. A stunning animated 3-D '
     'background (Vanta HALO) loads immediately. A status badge shows whether '
     'the backend server is awake (yellow = waking, green = ready).')

heading('Step 2 — Select a Filter Mode', 2)
para('Four filter buttons are shown:\n'
     '  ✏️  Pencil Sketch\n'
     '  ✨  Anime Art\n'
     '  🎨  Painting\n'
     '  ✂️  Remove Background\n\n'
     'The user clicks the desired filter. The active button is highlighted.')

heading('Step 3 — Upload a Photo', 2)
para('The user clicks the dashed upload zone (📸 "Click to Upload Photo") and '
     'selects a JPG or PNG from their device. The original image immediately '
     'previews on screen.')

heading('Step 4 — Apply the Filter', 2)
para('The user clicks the orange "Apply … Effect" button. The frontend sends '
     'the image file to the FastAPI backend via an HTTP POST request. A spinner '
     'animation ("Applying magic filters…") is shown while processing.')

heading('Step 5 — View and Download the Result', 2)
para('When the backend returns the processed image, it is displayed side-by-side '
     'with the original. A "Download Result" link lets the user save the output '
     'as JPG (or PNG for background removal).')

# ══════════════════════════════════════════════════════════════════════════════
#  6. BACKEND CODE — main.py
# ══════════════════════════════════════════════════════════════════════════════
heading('6. Backend Code — main.py', 1)
para('The entire backend is a single FastAPI file. It exposes four POST endpoints:')

table_2col([
    ('/api/sketch',     'Converts image to pencil sketch using grayscale + Gaussian blur + color dodge'),
    ('/api/anime',      'Applies anime/cartoon style using edge-preserving filter + adaptive thresholding'),
    ('/api/painting',   'Creates an oil-painting effect using edge-preserving filter + saturation boost'),
    ('/api/remove-bg',  'Removes image background using the rembg AI model (u2netp)'),
])
doc.add_paragraph()

heading('6.1 Full Backend Code', 2)
code_block(
"""import cv2
import numpy as np
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import Response
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI()

# rembg is lazy-loaded only when remove-bg is called
# to avoid OOM on startup (u2netp model = 4.7MB vs u2net = 176MB)
_rembg_session = None

def get_rembg_session():
    global _rembg_session
    if _rembg_session is None:
        from rembg import new_session
        _rembg_session = new_session("u2netp")
    return _rembg_session

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.get("/")
def read_root():
    return {"message": "Image Editing Backend"}

# ── Pencil Sketch ──────────────────────────────────────────────────────────
@app.post("/api/sketch")
async def convert_to_sketch(file: UploadFile = File(...)):
    contents = await file.read()
    nparr = np.frombuffer(contents, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

    gray_image    = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    inverted      = cv2.bitwise_not(gray_image)
    blurred       = cv2.GaussianBlur(inverted, (21, 21), 0)
    inv_blur      = cv2.bitwise_not(blurred)
    sketch        = cv2.divide(gray_image, inv_blur, scale=256.0)

    # Gamma correction for richer pencil strokes
    gamma   = 0.6
    table   = np.array([((i/255.0)**(1/gamma))*255
                        for i in range(256)]).astype("uint8")
    sketch  = cv2.LUT(sketch, table)

    _, enc = cv2.imencode('.jpg', sketch)
    return Response(content=enc.tobytes(), media_type="image/jpeg")

# ── Anime Art ──────────────────────────────────────────────────────────────
@app.post("/api/anime")
async def apply_anime(file: UploadFile = File(...)):
    contents = await file.read()
    nparr = np.frombuffer(contents, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

    h, w = img.shape[:2]
    max_dim = 1000
    scale   = 1.0
    if max(h, w) > max_dim:
        scale = max_dim / max(h, w)
        img   = cv2.resize(img, (int(w*scale), int(h*scale)))

    color = cv2.edgePreservingFilter(img, flags=1, sigma_s=50, sigma_r=0.4)

    hsv = cv2.cvtColor(color, cv2.COLOR_BGR2HSV).astype("float32")
    hh, ss, vv = cv2.split(hsv)
    ss = np.clip(ss * 1.25, 0, 255)
    vv = np.clip(vv * 1.05, 0, 255)
    color = cv2.cvtColor(cv2.merge([hh, ss, vv]).astype("uint8"), cv2.COLOR_HSV2BGR)

    gray  = cv2.medianBlur(cv2.cvtColor(img, cv2.COLOR_BGR2GRAY), 5)
    edges = cv2.adaptiveThreshold(gray, 255,
                cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 9, 9)
    anime = cv2.bitwise_and(color, color, mask=edges)

    if scale != 1.0:
        anime = cv2.resize(anime, (w, h))

    _, enc = cv2.imencode('.jpg', anime)
    return Response(content=enc.tobytes(), media_type="image/jpeg")

# ── Oil Painting ───────────────────────────────────────────────────────────
@app.post("/api/painting")
async def apply_painting(file: UploadFile = File(...)):
    contents = await file.read()
    nparr = np.frombuffer(contents, np.uint8)
    img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

    h, w   = img.shape[:2]
    max_dim = 1000
    scale   = 1.0
    if max(h, w) > max_dim:
        scale = max_dim / max(h, w)
        img   = cv2.resize(img, (int(w*scale), int(h*scale)))

    painted = cv2.edgePreservingFilter(img, flags=1, sigma_s=60, sigma_r=0.4)

    hsv = cv2.cvtColor(painted, cv2.COLOR_BGR2HSV).astype("float32")
    hc, s, v = cv2.split(hsv)
    s = np.clip(s * 1.2, 0, 255)
    painted = cv2.cvtColor(cv2.merge([hc, s, v]).astype("uint8"), cv2.COLOR_HSV2BGR)

    kernel  = np.array([[0,-0.5,0],[-0.5,3,-0.5],[0,-0.5,0]])
    painted = cv2.filter2D(painted, -1, kernel)

    if scale != 1.0:
        painted = cv2.resize(painted, (w, h))

    _, enc = cv2.imencode('.jpg', painted)
    return Response(content=enc.tobytes(), media_type="image/jpeg")

# ── Background Removal ─────────────────────────────────────────────────────
@app.post("/api/remove-bg")
async def remove_background(file: UploadFile = File(...)):
    from rembg import remove
    contents = await file.read()
    session  = get_rembg_session()
    output   = remove(contents, session=session)
    return Response(content=output, media_type="image/png")
"""
)

# ══════════════════════════════════════════════════════════════════════════════
#  7. FRONTEND CODE — App.jsx
# ══════════════════════════════════════════════════════════════════════════════
heading('7. Frontend Code — App.jsx', 1)
para(
    'The frontend is a single React component (App.jsx). Key responsibilities:\n'
    '  • Poll the backend every 5 s on startup to detect when it is awake\n'
    '  • Display a server-status badge (yellow = waking, green = ready)\n'
    '  • Handle file upload and send to the correct API endpoint\n'
    '  • Show original + result images side-by-side\n'
    '  • Provide a download link for the processed image'
)

heading('7.1 Key State Variables', 2)
table_2col([
    ('originalImage',  'Object URL of the uploaded image shown in the browser'),
    ('resultImage',    'Object URL of the processed image returned by the backend'),
    ('isLoading',      'Boolean — true while waiting for backend response'),
    ('activeMode',     'Currently selected filter: sketch / anime / painting / remove-bg'),
    ('serverStatus',   '"waking" or "ready" — controls the status badge color'),
])
doc.add_paragraph()

heading('7.2 API Communication Logic', 2)
code_block(
"""const processImage = async () => {
  setIsLoading(true);
  const formData = new FormData();
  formData.append('file', fileInputRef.current.files[0]);

  let endpoint = '/api/sketch';
  if (activeMode === 'remove-bg') endpoint = '/api/remove-bg';
  if (activeMode === 'anime')     endpoint = '/api/anime';
  if (activeMode === 'painting')  endpoint = '/api/painting';

  const response = await fetch(`${BACKEND_URL}${endpoint}`, {
    method: 'POST',
    body: formData,
  });

  if (response.ok) {
    const blob = await response.blob();
    setResultImage(URL.createObjectURL(blob));
  }
  setIsLoading(false);
};
"""
)

# ══════════════════════════════════════════════════════════════════════════════
#  8. IMAGE FILTER ALGORITHMS — DETAILED EXPLANATION
# ══════════════════════════════════════════════════════════════════════════════
heading('8. Image Filter Algorithms — Detailed Explanation', 1)

heading('8.1 Pencil Sketch', 2)
para(
    'Algorithm steps:\n'
    '  1. Convert the colour image to Grayscale.\n'
    '  2. Invert the grayscale image (bitwise NOT).\n'
    '  3. Apply Gaussian Blur (kernel 21×21) to the inverted image.\n'
    '  4. Invert the blurred image again.\n'
    '  5. Use Color Dodge blend: divide grayscale by inverted-blur (scale=256).\n'
    '  6. Apply Gamma Correction (γ=0.6) using a LUT to darken and enrich strokes.\n\n'
    'Result: A realistic pencil-sketch with dark, expressive strokes.'
)

heading('8.2 Anime Art', 2)
para(
    'Algorithm steps:\n'
    '  1. Resize large images to max 1000 px (performance optimisation).\n'
    '  2. Apply Edge-Preserving Filter (sigma_s=50, sigma_r=0.4) for cel-shaded colours.\n'
    '  3. Boost saturation ×1.25 and brightness ×1.05 in HSV colour space.\n'
    '  4. Apply Median Blur on the grayscale version to remove noise.\n'
    '  5. Extract fine outlines via Adaptive Threshold (block=9, C=9).\n'
    '  6. Combine coloured image with edge mask using bitwise AND.\n\n'
    'Result: A vibrant anime/cartoon appearance with clean black outlines.'
)

heading('8.3 Oil Painting', 2)
para(
    'Algorithm steps:\n'
    '  1. Resize large images to max 1000 px.\n'
    '  2. Apply Edge-Preserving Filter (sigma_s=60, sigma_r=0.4) — heavier smoothing than anime.\n'
    '  3. Boost saturation ×1.2 in HSV colour space.\n'
    '  4. Apply a custom Sharpening Kernel (3×3) to simulate brush-stroke texture.\n\n'
    'Result: A rich oil-painting / watercolour appearance.'
)

heading('8.4 Background Removal', 2)
para(
    'Algorithm steps:\n'
    '  1. Lazy-load the rembg session with the lightweight u2netp model (4.7 MB) '
    'only on the first background-removal request.\n'
    '  2. Pass the raw image bytes to rembg.remove().\n'
    '  3. Return the result as a PNG with a transparent background.\n\n'
    'Result: Subject is cleanly isolated on a transparent background. '
    'The frontend renders a checkerboard pattern behind the PNG to show transparency.'
)

# ══════════════════════════════════════════════════════════════════════════════
#  9. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
heading('9. Deployment', 1)

heading('9.1 Backend — Render.com', 2)
para(
    'The FastAPI backend is deployed on Render\'s free tier. '
    'The render.yaml file configures the service:'
)
code_block(
"""services:
  - type: web
    name: photo-converter-backend
    runtime: python
    buildCommand: pip install -r requirements.txt
    startCommand: uvicorn main:app --host 0.0.0.0 --port $PORT
    envVars:
      - key: PYTHON_VERSION
        value: 3.11.0
"""
)
para('⚠️  Note: Render\'s free tier spins down after 15 minutes of inactivity. '
     'The frontend polls the backend every 5 seconds for up to 2 minutes to '
     'handle the "cold start" gracefully.')

heading('9.2 Frontend — Vercel', 2)
para('The React app is deployed on Vercel. The vercel.json configures the build:')
code_block(
"""{
  "buildCommand": "npm run build",
  "outputDirectory": "dist",
  "framework": "vite"
}
"""
)
para('The environment variable VITE_BACKEND_URL is set in Vercel\'s dashboard to '
     'point to the Render backend URL, allowing cross-device access.')

# ══════════════════════════════════════════════════════════════════════════════
#  10. REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
heading('10. Python Dependencies (requirements.txt)', 1)
table_2col([
    ('fastapi==0.115.0',                       'Web framework for the API'),
    ('uvicorn==0.30.6',                        'ASGI server to run FastAPI'),
    ('python-multipart==0.0.9',               'Required to receive file uploads'),
    ('numpy==1.26.4',                          'Array operations for image data'),
    ('opencv-contrib-python-headless==4.9.0.80', 'Computer vision (no GUI needed)'),
    ('pillow==10.4.0',                         'Image I/O support for rembg'),
    ('rembg==2.0.57',                          'AI background removal library'),
    ('onnxruntime==1.16.3',                    'Runs the u2netp neural network model'),
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  11. OUTPUT SCREENSHOTS
# ══════════════════════════════════════════════════════════════════════════════
heading('11. Sample Output Screenshots', 1)
para('The following images show actual outputs generated by the application filters:')

base = r"c:\Users\Shama\.gemini\antigravity\scratch\photo-converter"
samples = [
    ("test_image.jpg",  "Original Input Image"),
    ("test_out_A.jpg",  "Output A — Anime Art Filter"),
    ("test_out_B.jpg",  "Output B — Anime Art Filter (variant)"),
    ("test_out_C.jpg",  "Output C — Pencil Sketch Filter"),
    ("test_out_D.jpg",  "Output D — Oil Painting Filter"),
]

for fname, caption in samples:
    path = os.path.join(base, fname)
    add_image_if_exists(path, caption, width=Inches(4.5))
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  12. HOW TO RUN LOCALLY
# ══════════════════════════════════════════════════════════════════════════════
heading('12. How to Run Locally', 1)

heading('Method A — Double-Click start.bat (Windows)', 2)
para('Simply double-click the start.bat file in the project root. It automatically:\n'
     '  1. Starts the FastAPI backend on http://localhost:5000\n'
     '  2. Starts the React frontend on http://localhost:5173\n'
     '  3. Opens the browser automatically')

heading('Method B — Manual Steps', 2)
para('Backend:')
code_block(
"""cd backend
pip install -r requirements.txt
uvicorn main:app --host 0.0.0.0 --port 5000 --reload
"""
)
para('Frontend:')
code_block(
"""cd frontend
npm install
npm run dev
"""
)
para('Then open http://localhost:5173 in your browser.')

# ══════════════════════════════════════════════════════════════════════════════
#  13. KEY FEATURES SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading('13. Key Features Summary', 1)
table_2col([
    ('4 Art Filters',           'Sketch, Anime, Painting, Background Removal'),
    ('Real-time Preview',       'Original and result shown side-by-side'),
    ('Download Button',         'One-click download of processed image'),
    ('Server Status Badge',     'Visual indicator showing if backend is awake'),
    ('Animated Background',     'Vanta.js HALO 3-D animated gradient background'),
    ('Responsive Design',       'Works on mobile, tablet, and desktop'),
    ('Memory Optimised',        'u2netp model (4.7 MB) instead of u2net (176 MB)'),
    ('Auto Image Resize',       'Images >1000px are resized before processing to avoid OOM'),
    ('Cross-Device Access',     'Deployed on Vercel + Render, accessible worldwide'),
    ('CORS Enabled',            'Backend accepts requests from any frontend origin'),
])
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = os.path.join(base, "Photo_Converter_Project_Report.docx")
doc.save(out_path)
print("Word document saved to: " + out_path)
